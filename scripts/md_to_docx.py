from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document


@dataclass
class Options:
    input_md: Path
    output_docx: Path
    title: str | None = None


def _is_table_divider(line: str) -> bool:
    s = line.strip()
    if "|" not in s:
        return False
    # A markdown divider row like: |---|---:|---|
    return bool(re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", s))


def _split_table_row(line: str) -> list[str]:
    # Trim outer pipes then split; allow escaped pipes is out of scope for this lightweight exporter.
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def _add_heading(doc: Document, text: str, level: int) -> None:
    level = max(1, min(4, level))
    doc.add_heading(text.strip(), level=level)


def _add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _add_bullet(doc: Document, text: str, level: int = 0) -> None:
    # Word has built-in "List Bullet" styles; indentation is limited without custom numbering definitions.
    p = doc.add_paragraph(text.strip(), style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = p.paragraph_format.left_indent  # keep default; placeholder


def _add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text.strip(), style="List Number")


def _add_code_block(doc: Document, lines: Iterable[str]) -> None:
    # Use monospace-ish by using "No Spacing"; better formatting would require custom styles.
    p = doc.add_paragraph("", style="No Spacing")
    run = p.add_run("\n".join(lines).rstrip("\n"))
    # Best-effort: keep code readable
    run.font.name = "Consolas"


def export_md_to_docx(opts: Options) -> None:
    md = opts.input_md.read_text(encoding="utf-8").splitlines()
    doc = Document()

    if opts.title:
        doc.add_heading(opts.title, level=0)

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(md):
        line = md[i]

        # Code fences
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                _add_code_block(doc, code_buf)
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Tables (very basic)
        if "|" in line and i + 1 < len(md) and _is_table_divider(md[i + 1]):
            header = _split_table_row(line)
            i += 2  # skip divider
            rows: list[list[str]] = []
            while i < len(md) and "|" in md[i] and md[i].strip():
                rows.append(_split_table_row(md[i]))
                i += 1

            cols = max([len(header)] + [len(r) for r in rows] + [1])
            table = doc.add_table(rows=1 + len(rows), cols=cols)
            table.style = "Table Grid"

            for c in range(cols):
                table.cell(0, c).text = header[c] if c < len(header) else ""
                for r_idx, row in enumerate(rows, start=1):
                    table.cell(r_idx, c).text = row[c] if c < len(row) else ""
            continue

        s = line.rstrip()
        if not s.strip():
            # blank line → spacing
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            hashes, text = m.group(1), m.group(2)
            level = len(hashes)
            # Map md levels to word heading levels (1..4)
            mapped = 1 if level == 1 else 2 if level == 2 else 3 if level == 3 else 4
            _add_heading(doc, text, mapped)
            i += 1
            continue

        # Bullets
        if re.match(r"^\s*-\s+", s):
            _add_bullet(doc, re.sub(r"^\s*-\s+", "", s))
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", s):
            _add_numbered(doc, re.sub(r"^\s*\d+\.\s+", "", s))
            i += 1
            continue

        # Default paragraph
        _add_paragraph(doc, s)
        i += 1

    opts.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(opts.output_docx))


def main(argv: list[str]) -> int:
    if len(argv) < 3:
        print("Usage: python scripts/md_to_docx.py <input.md> <output.docx>")
        return 2

    input_md = Path(argv[1]).resolve()
    output_docx = Path(argv[2]).resolve()
    if not input_md.exists():
        print(f"Input file not found: {input_md}")
        return 2

    export_md_to_docx(Options(input_md=input_md, output_docx=output_docx, title=None))
    print(f"Wrote: {output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))


